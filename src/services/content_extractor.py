"""
Content Extractor for AWS Content Monitor.
"""
import hashlib
import re
from datetime import datetime
from typing import List, Dict, Any, Optional
from urllib.parse import urljoin, urlparse
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
import PyPDF2
from io import BytesIO
import docx
import tempfile
import os

from ..models.base import ExtractedContent, ContentSection, ValidationError
from ..models.enums import SourceType


class ContentExtractor:
    """Converts diverse document formats into normalized, comparable text while preserving structure."""
    
    def __init__(self, request_timeout: int = 30):
        """Initialize the content extractor."""
        self.request_timeout = request_timeout
        
        # Common navigation and non-content selectors
        self.navigation_selectors = [
            'nav', 'header', 'footer', '.navigation', '.nav', '.menu',
            '.sidebar', '.breadcrumb', '.pagination', '.social-share',
            '.advertisement', '.ads', '.banner', '.popup', '.modal'
        ]
        
        # Content selectors (in order of preference)
        self.content_selectors = [
            'main', 'article', '.content', '.main-content', '.post-content',
            '.entry-content', '.article-content', '.page-content', '#content'
        ]
    
    def _make_request(self, url: str) -> requests.Response:
        """Make HTTP request to fetch content."""
        headers = {
            'User-Agent': 'AWS-Content-Monitor/1.0 (https://github.com/aws-content-monitor)'
        }
        
        response = requests.get(url, headers=headers, timeout=self.request_timeout)
        response.raise_for_status()
        return response
    
    def extract_html(self, url: str) -> ExtractedContent:
        """Extract content from HTML web page."""
        try:
            response = self._make_request(url)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove navigation and non-content elements
            self._remove_navigation_elements(soup)
            
            # Find main content area
            content_area = self._find_main_content(soup)
            
            # Extract sections with hierarchy
            sections = self._extract_sections(content_area, url)
            
            # Generate content hash
            content_hash = self._calculate_content_hash(sections)
            
            # Create extraction metadata
            metadata = {
                'url': url,
                'title': self._extract_title(soup),
                'extraction_method': 'html_parser',
                'content_length': sum(len(s.content) for s in sections),
                'section_count': len(sections),
                'response_status': response.status_code,
                'content_type': response.headers.get('content-type', ''),
                'last_modified': response.headers.get('last-modified', ''),
                'etag': response.headers.get('etag', '')
            }
            
            return ExtractedContent(
                source_id=self._generate_source_id(url),
                content_hash=content_hash,
                sections=sections,
                extracted_at=datetime.utcnow(),
                extraction_metadata=metadata
            )
            
        except Exception as e:
            raise ValidationError(f"Failed to extract HTML content from {url}: {str(e)}", "url")
    
    def _remove_navigation_elements(self, soup: BeautifulSoup) -> None:
        """Remove navigation, headers, footers, and other non-content elements."""
        for selector in self.navigation_selectors:
            for element in soup.select(selector):
                element.decompose()
        
        # Remove script and style tags
        for tag in soup(['script', 'style', 'noscript']):
            tag.decompose()
        
        # Remove comments
        for comment in soup.find_all(string=lambda text: isinstance(text, str) and text.strip().startswith('<!--')):
            comment.extract()
    
    def _find_main_content(self, soup: BeautifulSoup) -> Tag:
        """Find the main content area of the page."""
        # Try content selectors in order of preference
        for selector in self.content_selectors:
            content_area = soup.select_one(selector)
            if content_area and self._has_substantial_content(content_area):
                return content_area
        
        # Fallback to body if no specific content area found
        body = soup.find('body')
        if body:
            return body
        
        # Last resort: use the entire soup
        return soup
    
    def _has_substantial_content(self, element: Tag) -> bool:
        """Check if an element has substantial text content."""
        text = element.get_text(strip=True)
        return len(text) > 100  # Minimum content threshold
    
    def _extract_sections(self, content_area: Tag, base_url: str) -> List[ContentSection]:
        """Extract sections with hierarchical structure."""
        sections = []
        current_section = None
        section_counter = 0
        
        # Find all heading tags and content between them
        elements = content_area.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'section', 'article'])
        
        for element in elements:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Start new section
                if current_section and current_section['content'].strip():
                    sections.append(self._create_content_section(current_section, section_counter))
                    section_counter += 1
                
                level = int(element.name[1])  # Extract number from h1, h2, etc.
                current_section = {
                    'heading': element.get_text(strip=True),
                    'content': '',
                    'level': level,
                    'position': section_counter
                }
            
            elif current_section is not None:
                # Add content to current section
                text = self._extract_clean_text(element)
                if text.strip():
                    current_section['content'] += text + '\n'
            
            else:
                # Content before first heading - create default section
                text = self._extract_clean_text(element)
                if text.strip():
                    if current_section is None:
                        current_section = {
                            'heading': 'Introduction',
                            'content': '',
                            'level': 1,
                            'position': 0
                        }
                    current_section['content'] += text + '\n'
        
        # Add final section
        if current_section and current_section['content'].strip():
            sections.append(self._create_content_section(current_section, section_counter))
        
        # If no sections found, create a single section with all content
        if not sections:
            all_text = content_area.get_text(separator='\n', strip=True)
            if all_text:
                sections.append(ContentSection(
                    id='section_0',
                    heading='Content',
                    content=self.normalize_text(all_text),
                    level=1,
                    position=0
                ))
        
        return sections
    
    def _create_content_section(self, section_data: Dict, position: int) -> ContentSection:
        """Create a ContentSection from extracted data."""
        return ContentSection(
            id=f"section_{position}",
            heading=section_data['heading'],
            content=self.normalize_text(section_data['content']),
            level=section_data['level'],
            position=position
        )
    
    def _extract_clean_text(self, element: Tag) -> str:
        """Extract clean text from an element."""
        if not element:
            return ''
        
        # Skip if element is likely navigation or metadata
        if element.get('class'):
            classes = ' '.join(element.get('class', []))
            if any(nav_class in classes.lower() for nav_class in ['nav', 'menu', 'sidebar', 'footer', 'header']):
                return ''
        
        # Get text with some formatting preserved
        text = element.get_text(separator=' ', strip=True)
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text
    
    def _extract_title(self, soup: BeautifulSoup) -> str:
        """Extract page title."""
        title_tag = soup.find('title')
        if title_tag:
            return title_tag.get_text(strip=True)
        
        # Fallback to first h1
        h1 = soup.find('h1')
        if h1:
            return h1.get_text(strip=True)
        
        return 'Untitled'
    
    def normalize_text(self, text: str) -> str:
        """Normalize text for consistent comparisons."""
        if not text:
            return ''
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        # Normalize quotes
        text = re.sub(r'[""''`]', '"', text)
        
        # Normalize dashes
        text = re.sub(r'[–—]', '-', text)
        
        # Remove zero-width characters
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)
        
        # Normalize line endings
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Normalize unicode characters
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
        
        return text
    
    def validate_extracted_content(self, content: ExtractedContent) -> Dict[str, Any]:
        """Validate extracted content completeness and quality."""
        validation_result = {
            'is_valid': True,
            'warnings': [],
            'errors': [],
            'metrics': {}
        }
        
        try:
            # Check basic structure
            if not content.sections:
                validation_result['errors'].append("No sections found in extracted content")
                validation_result['is_valid'] = False
            
            # Calculate content metrics
            total_content_length = sum(len(section.content) for section in content.sections)
            non_empty_sections = len([s for s in content.sections if s.content.strip()])
            
            validation_result['metrics'] = {
                'total_sections': len(content.sections),
                'non_empty_sections': non_empty_sections,
                'total_content_length': total_content_length,
                'average_section_length': total_content_length / len(content.sections) if content.sections else 0,
                'content_hash': content.content_hash,
                'extraction_method': content.extraction_metadata.get('extraction_method', 'unknown')
            }
            
            # Quality checks
            if total_content_length < 100:
                validation_result['warnings'].append("Very short content extracted (< 100 characters)")
            
            if non_empty_sections == 0:
                validation_result['errors'].append("No sections contain actual content")
                validation_result['is_valid'] = False
            
            # Check for duplicate sections
            section_hashes = set()
            for section in content.sections:
                section_hash = hashlib.md5(section.content.encode()).hexdigest()
                if section_hash in section_hashes:
                    validation_result['warnings'].append(f"Duplicate content detected in section: {section.heading}")
                section_hashes.add(section_hash)
            
            # Check section hierarchy
            levels = [section.level for section in content.sections]
            if levels and (min(levels) < 1 or max(levels) > 6):
                validation_result['warnings'].append("Section levels outside normal range (1-6)")
            
            # Validate content hash
            calculated_hash = self._calculate_content_hash(content.sections)
            if calculated_hash != content.content_hash:
                validation_result['errors'].append("Content hash mismatch - content may be corrupted")
                validation_result['is_valid'] = False
            
        except Exception as e:
            validation_result['errors'].append(f"Validation failed: {str(e)}")
            validation_result['is_valid'] = False
        
        return validation_result
    
    def generate_deterministic_output(self, content: ExtractedContent) -> str:
        """Generate deterministic output suitable for comparison."""
        # Sort sections by position to ensure consistent ordering
        sorted_sections = sorted(content.sections, key=lambda s: s.position)
        
        output_parts = []
        
        for section in sorted_sections:
            # Normalize section heading and content
            normalized_heading = self.normalize_text(section.heading)
            normalized_content = self.normalize_text(section.content)
            
            # Create deterministic section representation
            section_repr = f"[SECTION:{section.level}:{normalized_heading}]\n{normalized_content}\n[/SECTION]\n"
            output_parts.append(section_repr)
        
        return ''.join(output_parts)
    
    def compare_content_similarity(self, content1: ExtractedContent, content2: ExtractedContent) -> Dict[str, Any]:
        """Compare two extracted contents for similarity."""
        # Generate deterministic representations
        repr1 = self.generate_deterministic_output(content1)
        repr2 = self.generate_deterministic_output(content2)
        
        # Calculate similarity metrics
        similarity_result = {
            'identical': repr1 == repr2,
            'hash_match': content1.content_hash == content2.content_hash,
            'section_count_match': len(content1.sections) == len(content2.sections),
            'content_length_diff': abs(len(repr1) - len(repr2)),
            'similarity_ratio': 0.0
        }
        
        # Calculate text similarity using simple ratio
        if repr1 and repr2:
            # Simple character-based similarity
            max_len = max(len(repr1), len(repr2))
            min_len = min(len(repr1), len(repr2))
            
            if max_len > 0:
                # Count matching characters at the beginning
                matching_chars = 0
                for i in range(min_len):
                    if repr1[i] == repr2[i]:
                        matching_chars += 1
                    else:
                        break
                
                similarity_result['similarity_ratio'] = matching_chars / max_len
        
        return similarity_result
    
    def extract_content_statistics(self, content: ExtractedContent) -> Dict[str, Any]:
        """Extract statistical information about the content."""
        stats = {
            'section_count': len(content.sections),
            'total_characters': 0,
            'total_words': 0,
            'total_sentences': 0,
            'heading_distribution': {},
            'average_section_length': 0,
            'longest_section': 0,
            'shortest_section': float('inf'),
            'content_density': 0.0
        }
        
        section_lengths = []
        
        for section in content.sections:
            content_text = section.content
            char_count = len(content_text)
            word_count = len(content_text.split())
            sentence_count = len(re.split(r'[.!?]+', content_text))
            
            stats['total_characters'] += char_count
            stats['total_words'] += word_count
            stats['total_sentences'] += sentence_count
            
            section_lengths.append(char_count)
            
            # Track heading levels
            level_key = f'level_{section.level}'
            stats['heading_distribution'][level_key] = stats['heading_distribution'].get(level_key, 0) + 1
            
            # Update min/max section lengths
            stats['longest_section'] = max(stats['longest_section'], char_count)
            if char_count > 0:  # Only consider non-empty sections for shortest
                stats['shortest_section'] = min(stats['shortest_section'], char_count)
        
        # Calculate averages
        if section_lengths:
            stats['average_section_length'] = sum(section_lengths) / len(section_lengths)
            
            # Content density: ratio of content to structure
            structure_overhead = len(content.sections) * 50  # Estimate overhead per section
            stats['content_density'] = stats['total_characters'] / (stats['total_characters'] + structure_overhead)
        
        # Handle edge case where no content was found
        if stats['shortest_section'] == float('inf'):
            stats['shortest_section'] = 0
        
        return stats
    
    def _calculate_content_hash(self, sections: List[ContentSection]) -> str:
        """Calculate hash of content for change detection."""
        content_str = ''.join(f"{s.heading}:{s.content}" for s in sections)
        return hashlib.sha256(content_str.encode('utf-8')).hexdigest()
    
    def _generate_source_id(self, url: str) -> str:
        """Generate source ID from URL."""
        return hashlib.md5(url.encode('utf-8')).hexdigest()
    
    def preserve_structure(self, content: str) -> Dict[str, Any]:
        """Preserve document structure for analysis."""
        # This is a placeholder for more advanced structure preservation
        # Could include things like:
        # - Table structures
        # - List hierarchies
        # - Link relationships
        # - Image references
        
        structure = {
            'paragraphs': content.split('\n\n'),
            'sentences': re.split(r'[.!?]+', content),
            'word_count': len(content.split()),
            'character_count': len(content)
        }
        
        return structure
    
    def extract_pdf(self, file_path_or_url: str) -> ExtractedContent:
        """Extract content from PDF document."""
        try:
            # Handle both file paths and URLs
            if file_path_or_url.startswith(('http://', 'https://')):
                # Download PDF from URL
                response = self._make_request(file_path_or_url)
                pdf_content = BytesIO(response.content)
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            else:
                # Read from local file
                with open(file_path_or_url, 'rb') as f:
                    pdf_content = BytesIO(f.read())
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            
            # Extract text from PDF
            pdf_reader = PyPDF2.PdfReader(pdf_content)
            
            sections = []
            current_section = None
            section_counter = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if not page_text.strip():
                        continue
                    
                    # Try to identify sections by looking for headings
                    lines = page_text.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # Heuristic: lines that are short, uppercase, or end with certain patterns might be headings
                        if self._is_likely_heading(line):
                            # Start new section
                            if current_section and current_section['content'].strip():
                                sections.append(self._create_content_section(current_section, section_counter))
                                section_counter += 1
                            
                            current_section = {
                                'heading': line,
                                'content': '',
                                'level': self._estimate_heading_level(line),
                                'position': section_counter
                            }
                        else:
                            # Add to current section
                            if current_section is None:
                                current_section = {
                                    'heading': f'Page {page_num + 1}',
                                    'content': '',
                                    'level': 1,
                                    'position': section_counter
                                }
                            
                            current_section['content'] += line + ' '
                
                except Exception as e:
                    # Log error but continue with other pages
                    print(f"Warning: Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Add final section
            if current_section and current_section['content'].strip():
                sections.append(self._create_content_section(current_section, section_counter))
            
            # If no sections found, create a single section
            if not sections:
                all_text = ''
                for page in pdf_reader.pages:
                    try:
                        all_text += page.extract_text() + '\n'
                    except:
                        continue
                
                if all_text.strip():
                    sections.append(ContentSection(
                        id='section_0',
                        heading='Document Content',
                        content=self.normalize_text(all_text),
                        level=1,
                        position=0
                    ))
            
            # Generate content hash
            content_hash = self._calculate_content_hash(sections)
            
            # Create extraction metadata
            metadata = {
                'source': source_ref,
                'extraction_method': 'pdf_parser',
                'page_count': len(pdf_reader.pages),
                'content_length': sum(len(s.content) for s in sections),
                'section_count': len(sections),
                'pdf_metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else {}
            }
            
            return ExtractedContent(
                source_id=source_id,
                content_hash=content_hash,
                sections=sections,
                extracted_at=datetime.utcnow(),
                extraction_metadata=metadata
            )
            
        except Exception as e:
            raise ValidationError(f"Failed to extract PDF content from {file_path_or_url}: {str(e)}", "file_path")
    
    def extract_docx(self, file_path_or_url: str) -> ExtractedContent:
        """Extract content from DOCX document."""
        try:
            # Handle both file paths and URLs
            if file_path_or_url.startswith(('http://', 'https://')):
                # Download DOCX from URL
                response = self._make_request(file_path_or_url)
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(response.content)
                    temp_path = temp_file.name
                
                try:
                    doc = docx.Document(temp_path)
                finally:
                    os.unlink(temp_path)  # Clean up temp file
                
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            else:
                # Read from local file
                doc = docx.Document(file_path_or_url)
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            
            sections = []
            current_section = None
            section_counter = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if paragraph is a heading based on style
                if paragraph.style.name.startswith('Heading') or self._is_likely_heading(text):
                    # Start new section
                    if current_section and current_section['content'].strip():
                        sections.append(self._create_content_section(current_section, section_counter))
                        section_counter += 1
                    
                    # Extract heading level from style name
                    level = 1
                    if paragraph.style.name.startswith('Heading'):
                        try:
                            level = int(paragraph.style.name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1
                    else:
                        level = self._estimate_heading_level(text)
                    
                    current_section = {
                        'heading': text,
                        'content': '',
                        'level': level,
                        'position': section_counter
                    }
                else:
                    # Add to current section
                    if current_section is None:
                        current_section = {
                            'heading': 'Document Content',
                            'content': '',
                            'level': 1,
                            'position': section_counter
                        }
                    
                    current_section['content'] += text + '\n'
            
            # Add final section
            if current_section and current_section['content'].strip():
                sections.append(self._create_content_section(current_section, section_counter))
            
            # If no sections found, create a single section
            if not sections:
                all_text = '\n'.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
                if all_text.strip():
                    sections.append(ContentSection(
                        id='section_0',
                        heading='Document Content',
                        content=self.normalize_text(all_text),
                        level=1,
                        position=0
                    ))
            
            # Generate content hash
            content_hash = self._calculate_content_hash(sections)
            
            # Create extraction metadata
            metadata = {
                'source': source_ref,
                'extraction_method': 'docx_parser',
                'paragraph_count': len(doc.paragraphs),
                'content_length': sum(len(s.content) for s in sections),
                'section_count': len(sections),
                'document_properties': {
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                }
            }
            
            return ExtractedContent(
                source_id=source_id,
                content_hash=content_hash,
                sections=sections,
                extracted_at=datetime.utcnow(),
                extraction_metadata=metadata
            )
            
        except Exception as e:
            raise ValidationError(f"Failed to extract DOCX content from {file_path_or_url}: {str(e)}", "file_path")
    
    def extract_text_file(self, file_path_or_url: str) -> ExtractedContent:
        """Extract content from plain text file."""
        try:
            # Handle both file paths and URLs
            if file_path_or_url.startswith(('http://', 'https://')):
                # Download text from URL
                response = self._make_request(file_path_or_url)
                text_content = response.text
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            else:
                # Read from local file
                with open(file_path_or_url, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                source_id = self._generate_source_id(file_path_or_url)
                source_ref = file_path_or_url
            
            # Split into sections based on double line breaks or other patterns
            sections = []
            
            # Try to identify sections by looking for patterns
            paragraphs = text_content.split('\n\n')
            current_section = None
            section_counter = 0
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                lines = paragraph.split('\n')
                first_line = lines[0].strip()
                
                # Check if first line looks like a heading
                if self._is_likely_heading(first_line) and len(lines) > 1:
                    # Start new section
                    if current_section and current_section['content'].strip():
                        sections.append(self._create_content_section(current_section, section_counter))
                        section_counter += 1
                    
                    current_section = {
                        'heading': first_line,
                        'content': '\n'.join(lines[1:]),
                        'level': self._estimate_heading_level(first_line),
                        'position': section_counter
                    }
                else:
                    # Add to current section
                    if current_section is None:
                        current_section = {
                            'heading': 'Text Content',
                            'content': '',
                            'level': 1,
                            'position': section_counter
                        }
                    
                    current_section['content'] += paragraph + '\n\n'
            
            # Add final section
            if current_section and current_section['content'].strip():
                sections.append(self._create_content_section(current_section, section_counter))
            
            # If no sections found, create a single section
            if not sections:
                sections.append(ContentSection(
                    id='section_0',
                    heading='Text Content',
                    content=self.normalize_text(text_content),
                    level=1,
                    position=0
                ))
            
            # Generate content hash
            content_hash = self._calculate_content_hash(sections)
            
            # Create extraction metadata
            metadata = {
                'source': source_ref,
                'extraction_method': 'text_parser',
                'content_length': len(text_content),
                'section_count': len(sections),
                'line_count': text_content.count('\n') + 1,
                'character_count': len(text_content)
            }
            
            return ExtractedContent(
                source_id=source_id,
                content_hash=content_hash,
                sections=sections,
                extracted_at=datetime.utcnow(),
                extraction_metadata=metadata
            )
            
        except Exception as e:
            raise ValidationError(f"Failed to extract text content from {file_path_or_url}: {str(e)}", "file_path")
    
    def _is_likely_heading(self, text: str) -> bool:
        """Heuristic to determine if text is likely a heading."""
        text = text.strip()
        
        # Empty text is not a heading
        if not text:
            return False
        
        # Very long text is unlikely to be a heading
        if len(text) > 100:
            return False
        
        # Check for common heading patterns
        heading_patterns = [
            r'^\d+\.?\s+',  # Numbered headings (1. or 1)
            r'^[A-Z][A-Z\s]+$',  # ALL CAPS
            r'^[A-Z][a-z\s]+:$',  # Title Case ending with colon
            r'^\*+\s+',  # Markdown-style headings
            r'^#+\s+',  # Markdown headings
        ]
        
        for pattern in heading_patterns:
            if re.match(pattern, text):
                return True
        
        # Check if text is short and doesn't end with punctuation (except colon)
        if len(text) < 50 and not text.endswith(('.', '!', '?', ';')) or text.endswith(':'):
            # Check if it's mostly title case
            words = text.split()
            if len(words) > 0:
                title_case_words = sum(1 for word in words if word[0].isupper())
                if title_case_words / len(words) > 0.5:
                    return True
        
        return False
    
    def _estimate_heading_level(self, text: str) -> int:
        """Estimate heading level based on text characteristics."""
        text = text.strip()
        
        # Check for numbered headings
        if re.match(r'^\d+\.', text):
            return 1
        elif re.match(r'^\d+\.\d+\.', text):
            return 2
        elif re.match(r'^\d+\.\d+\.\d+\.', text):
            return 3
        
        # Check for markdown-style headings
        if text.startswith('#'):
            return min(text.count('#'), 6)
        
        # Estimate based on length and case
        if len(text) < 20 and text.isupper():
            return 1
        elif len(text) < 40:
            return 2
        else:
            return 3
    
    def extract_content(self, source_url: str, source_type: SourceType) -> ExtractedContent:
        """Extract content based on source type."""
        if source_type == SourceType.HTML:
            return self.extract_html(source_url)
        elif source_type == SourceType.PDF:
            return self.extract_pdf(source_url)
        elif source_type == SourceType.DOCX:
            return self.extract_docx(source_url)
        elif source_type == SourceType.TXT:
            return self.extract_text_file(source_url)
        else:
            raise ValidationError(f"Unsupported source type: {source_type}", "source_type")